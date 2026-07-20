"""Document text extraction — PDF, DOCX, XLSX, ZIP, plain text.

Large portal PDFs (e.g. GeM) can extract as heavily duplicated token noise;
`clean_text` dedupes consecutive repeats before any LLM spend.
"""
from __future__ import annotations

import io
import logging
import re
import zipfile
from pathlib import Path

log = logging.getLogger("bidpilot.docparse")


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    # GeM bid PDFs often carry the technical-spec documents as files EMBEDDED
    # inside the main PDF — pull their text out too.
    try:
        for name, contents in (reader.attachments or {}).items():
            suffix = Path(name).suffix.lower()
            if suffix not in (".pdf", ".docx", ".xlsx", ".txt", ".csv"):
                continue
            blobs = contents if isinstance(contents, list) else [contents]
            for blob in blobs:
                inner = extract_bytes(bytes(blob), suffix)
                if inner.strip():
                    text += f"\n\n# Embedded file: {name}\n{inner}"
    except Exception as e:
        log.warning("embedded-attachment extraction failed: %s", e)
    return text


def _docx_text(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx_text(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_bytes(data: bytes, suffix: str) -> str:
    suffix = suffix.lower()
    try:
        if suffix == ".pdf":
            return _pdf_text(data)
        if suffix == ".docx":
            return _docx_text(data)
        if suffix in (".xlsx", ".xls"):
            return _xlsx_text(data)
        if suffix == ".zip":
            parts = []
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for name in zf.namelist():
                    inner = Path(name).suffix.lower()
                    if inner in (".pdf", ".docx", ".xlsx", ".txt", ".csv"):
                        parts.append(f"# File: {name}\n" + extract_bytes(zf.read(name), inner))
            return "\n\n".join(parts)
        return data.decode("utf-8", errors="replace")
    except Exception as e:  # a broken attachment must not kill the scan
        log.warning("failed to extract %s: %s", suffix, e)
        return ""


def extract_file(path: str | Path) -> str:
    p = Path(path)
    return extract_bytes(p.read_bytes(), p.suffix)


def clean_text(text: str, cap_chars: int = 40000) -> str:
    """Collapse whitespace, drop consecutive duplicate lines, cap length."""
    lines = []
    prev = None
    for line in text.splitlines():
        s = re.sub(r"[ \t]+", " ", line).strip()
        if not s or s == prev:
            prev = s
            continue
        lines.append(s)
        prev = s
    out = "\n".join(lines)
    # dedupe immediately repeated tokens ("Bid Bid Bid Number Number")
    out = re.sub(r"\b(\S+)(?: \1\b)+", r"\1", out)
    if len(out) > cap_chars:
        out = out[:cap_chars] + "\n[... truncated ...]"
    return out
